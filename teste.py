import time
import random
import subprocess
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document

# Assume-se que está no mesmo sitio que isto
executavel = "./a.out"
valoresnm = []
tempos = []
tabela_dados = []

def build_table(n):
    # Alocar tabela
    tab = [[0] * n for _ in range(n)]

    # Preencher tabela
    for i in range(n):
        for j in range(n):
            tab[i][j] = random.randint(1, n)

    return tab

def build_sequence(tab, n, m):
    # Alocar sequência
    seq = [0] * (m + 1)

    # Preencher sequência
    for i in range(m):
        seq[i] = random.randint(1, n)

    # Calcular uma saída válida eliminando vizinhos escolhidos aleatoriamente
    tmp = seq[:m]
    while len(tmp) > 1:
        i = random.randint(0, len(tmp) - 2)  # Escolher um índice aleatório
        tmp[i] = tab[tmp[i] - 1][tmp[i + 1] - 1]  # Atualizar tmp[i]
        tmp.pop(i + 1)  # Remover o próximo elemento

    # O último elemento do array seq é a resposta
    seq[m] = tmp[0]
    return seq

def create_test(n, m):
    tab = build_table(n)
    seq = build_sequence(tab, n, m)
    
    with open("test.in", "w") as f:
        f.write(f"{n} {m}\n")
        for linha in tab:
            f.write(" ".join(map(str, linha)) + "\n")
        f.write(" ".join(map(str, seq[:-1])) + "\n")
        f.write(str(seq[-1]) + "\n")


def run_tests():
    i = 1
    for n in range(5, 101, 5):
        for m in range(10, 1001, 15):
            # Gerar um teste válido
            create_test(n, m)
            
            # Correr os testes
            start = time.time()
            subprocess.run([executavel], stdin=open("test.in"), stdout=open("test.out", "w"))
            total = time.time() - start
            
            # Verifica se o tempo é aceitável
            if total <= 15:
                tempos.append(total)

                # Editar para melhor efeito
                fmn = (n ** 2) * (m ** 3)
                
                # Adicionar valores à lista
                valoresnm.append(fmn)
                tabela_dados.append({"n": n, "m": m, "f(n,m)": fmn, "tempo": total})

                print(f"Teste {i}: n={n}, m={m} -> Tempo de execução: {total:.4f} segundos")
            else:
                print(f"Teste {i}: n={n}, m={m} -> Tempo de execução: {total:.4f} segundos (excedeu 15 segundos e não será plotado)")
            
            i += 1

def calcular_valores_medios(num_entradas):
    if num_entradas > len(valoresnm):
        print("Número de entradas solicitado excede o tamanho das listas.")
        return

    # Inicializar a tabela de dados medianos
    tabela_dados_medianos = []
    indices = np.linspace(0, len(valoresnm)-1, num_entradas, dtype=int)

    # Percorrer a tabela de dados e adicionar indices medianos
    for i in indices:
        tabela_dados_medianos.append(tabela_dados[i])

    return tabela_dados_medianos

# Correr os testes
run_tests()

# Pergunta ao usuário quantas linhas deseja na tabela
num_linhas = int(input("Quantas linhas deseja na tabela? "))

# Cria um documento do Word e adiciona a tabela
doc = Document()
doc.add_heading('Resultados dos Testes', level=1)

# Adiciona a tabela ao documento
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'n'
hdr_cells[1].text = 'm'
hdr_cells[2].text = "f(n, m)"
hdr_cells[3].text = 'Tempos'

# Preenche a tabela com os dados
tabela_dados = calcular_valores_medios(num_linhas)
for row in tabela_dados:
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['n'])
    row_cells[1].text = str(row['m'])
    row_cells[2].text = f"{row['f(n,m)']:.4e}"
    row_cells[3].text = f"{row['tempo']:.4f}"

# Salva o documento
doc.save('resultados_testes.docx')

# Plota os pontos de dados originais (média dos tempos)
plt.figure(figsize=(15, 9))  # Tamanho da figura
plt.scatter(valoresnm, tempos, label="Dados experimentais", alpha=0.7, color="cyan", edgecolor='black', s=100)

# Ajusta uma curva polinomial de tendência para todos os dados combinados
degree = 1
coef = np.polyfit(valoresnm, tempos, degree)
print(f"Coeficientes da curva de tendência: y = {coef[0]}x + {coef[1]}")
poly_fn = np.poly1d(coef)
plt.plot(valoresnm, poly_fn(valoresnm), '--', label="Tendência global", color="red", linewidth=2)

# Personaliza o gráfico
plt.xlabel("f(n, m)", fontsize=14)
plt.ylabel("Tempo (s)", fontsize=14)
plt.title("Curva de Tendência para Tempo de Execução em Função de f(n, m)", fontsize=16)
plt.legend(fontsize=12)
plt.grid(True, linestyle='--', alpha=0.7)  # Adiciona uma grade
plt.xticks(fontsize=12)
plt.yticks(fontsize=12)

# Salva o gráfico como imagem
plt.savefig('grafico_resultados.png', dpi=300, bbox_inches='tight')

# Mostra o gráfico
plt.show()
